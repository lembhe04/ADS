"""
Document Q&A System using Streamlit, LangChain, LangGraph, and Hugging Face
A complete application for document-based question answering with chat history
and a self-correcting RAG workflow powered by LangGraph.
"""
#streamlit run c:/Users/ADMIN/Desktop/python/app.py

import streamlit as st
import os
import tempfile
import logging
from typing import List, Optional, Dict, Any, TypedDict
import hashlib
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Document processing imports
import PyPDF2
from docx import Document
import io

# LangChain and LangGraph imports
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.llms import HuggingFacePipeline
from langchain.schema import Document as LangChainDocument
from langchain_core.output_parsers import StrOutputParser
from langchain.prompts import PromptTemplate

# LangGraph imports
from langgraph.graph import END, StateGraph

# Transformers and sentence-transformers
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline
import torch

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load Hugging Face token from environment
HF_TOKEN = os.getenv('HUGGINGFACE_HUB_TOKEN') or os.getenv('HF_TOKEN')
if HF_TOKEN:
    os.environ['HUGGINGFACE_HUB_TOKEN'] = HF_TOKEN
    logger.info("Hugging Face token loaded from environment")
else:
    logger.warning("No Hugging Face token found in environment variables")

# Page configuration
st.set_page_config(
    page_title="Document Q&A System with LangGraph",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .chat-message {
        padding: 1rem;
        border-radius: 0.5rem;
        margin-bottom: 1rem;
        display: flex;
        flex-direction: column;
        color: black;
    }
    .user-message {
        background-color: #e3f2fd;
        border-left: 4px solid #2196f3;
        color: black;
    }
    .bot-message {
        background-color: #f1f8e9;
        border-left: 4px solid #4caf50;
        color: black;
    }
    .context-box {
        background-color: #fff3e0;
        border: 1px solid #ff9800;
        border-radius: 0.5rem;
        padding: 1rem;
        margin: 1rem 0;
        color: black;
    }
    .error-message {
        background-color: #ffebee;
        color: #c62828;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #f44336;
    }
</style>
""", unsafe_allow_html=True)


# --- Document Processing Classes ---

class DocumentProcessor:
    """Handles document processing and text extraction."""
    
    @staticmethod
    def extract_text_from_pdf(file) -> str:
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            return "".join(page.extract_text() for page in pdf_reader.pages)
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file) -> str:
        try:
            doc = Document(file)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(file) -> str:
        try:
            return str(file.read(), "utf-8")
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {e}")
            raise Exception(f"Failed to extract text from TXT: {str(e)}")

    @staticmethod
    def process_uploaded_file(uploaded_file) -> str:
        if uploaded_file is None: return ""
        file_extension = uploaded_file.name.split('.')[-1].lower()
        try:
            if file_extension == 'pdf': return DocumentProcessor.extract_text_from_pdf(uploaded_file)
            elif file_extension == 'docx': return DocumentProcessor.extract_text_from_docx(uploaded_file)
            elif file_extension == 'txt': return DocumentProcessor.extract_text_from_txt(uploaded_file)
            else: raise ValueError(f"Unsupported file type: {file_extension}")
        except Exception as e:
            logger.error(f"Error processing file {uploaded_file.name}: {e}")
            raise


class EmbeddingManager:
    """Manages embeddings and vector store operations."""
    
    def __init__(self):
        self.embeddings = None
        self.vector_store = None
    
    @st.cache_resource
    def load_embeddings(_self):
        try:
            model_kwargs = {'device': 'cuda' if torch.cuda.is_available() else 'cpu'}
            encode_kwargs = {'normalize_embeddings': True}
            if HF_TOKEN: model_kwargs['token'] = HF_TOKEN
            
            embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs=model_kwargs,
                encode_kwargs=encode_kwargs
            )
            logger.info("Embeddings loaded successfully")
            return embeddings
        except Exception as e:
            logger.error(f"Error loading embeddings: {e}")
            raise Exception(f"Failed to load embeddings: {str(e)}")
    
    def create_vector_store(self, documents: List[LangChainDocument]):
        try:
            if not self.embeddings: self.embeddings = self.load_embeddings()
            self.vector_store = FAISS.from_documents(documents=documents, embedding=self.embeddings)
            logger.info(f"Vector store created with {len(documents)} documents")
            return self.vector_store
        except Exception as e:
            logger.error(f"Error creating vector store: {e}")
            raise Exception(f"Failed to create vector store: {str(e)}")


# --- LangGraph RAG System ---

class GraphState(TypedDict):
    """Represents the state of our graph."""
    question: str
    documents: List[LangChainDocument]
    answer: str

class QASystem:
    """Manages the question-answering system using LangGraph."""
    
    def __init__(self):
        self.llm = None
        self.rag_graph = None
        self.retriever = None

    @st.cache_resource
    def load_llm(_self):
        """Load Hugging Face LLM with caching."""
        try:
            device = 0 if torch.cuda.is_available() else -1
            model_name = "google/flan-t5-large"
            load_kwargs = {'token': HF_TOKEN} if HF_TOKEN else {}
            
            tokenizer = AutoTokenizer.from_pretrained(model_name, **load_kwargs)
            model = AutoModelForSeq2SeqLM.from_pretrained(model_name, **load_kwargs)
            
            pipe = pipeline("text2text-generation", model=model, tokenizer=tokenizer, max_length=512, temperature=0.1, device=device)
            llm = HuggingFacePipeline(pipeline=pipe)
            logger.info("LLM loaded successfully")
            return llm
        except Exception as e:
            logger.error(f"Error loading LLM: {e}")
            raise Exception(f"Failed to load LLM: {str(e)}")

    # --- Graph Nodes ---
    def retrieve(self, state: GraphState) -> GraphState:
        """Retrieves documents from the vector store."""
        logger.info("Node: Retrieving documents")
        question = state['question']
        documents = self.retriever.invoke(question)
        return {"documents": documents, "question": question, "answer": ""}

    def grade_documents(self, state: GraphState) -> str:
        """Determines if the retrieved documents are relevant to the question."""
        logger.info("Node: Grading documents")
        question = state['question']
        documents = state['documents']
        
        prompt = PromptTemplate.from_template(
            """Given the user's question and the retrieved documents, determine if the documents are relevant to answer the question. 
            Answer with a single word, either 'yes' or 'no'.
            
            Question: {question}
            Documents: {documents}
            
            Relevant (yes/no):"""
        )
        
        grading_chain = prompt | self.llm | StrOutputParser()
        docs_content = "\n\n".join([d.page_content for d in documents])
        result = grading_chain.invoke({"question": question, "documents": docs_content})
        
        if "yes" in result.lower():
            logger.info("Decision: Documents are relevant, proceeding to generate.")
            return "generate"
        else:
            logger.info("Decision: Documents not relevant, using fallback.")
            return "fallback"

    def generate(self, state: GraphState) -> GraphState:
        """Generates an answer using the retrieved documents."""
        logger.info("Node: Generating answer")
        question = state['question']
        documents = state['documents']
        
        prompt = PromptTemplate.from_template(
            """You are an assistant for question-answering tasks. Use the following pieces of retrieved context to answer the question. 
            If you don't know the answer, just say that you don't know. Keep the answer concise.
            
            Question: {question}
            Context: {context}
            
            Answer:"""
        )
        
        rag_chain = prompt | self.llm | StrOutputParser()
        docs_content = "\n\n".join([d.page_content for d in documents])
        answer = rag_chain.invoke({"context": docs_content, "question": question})
        return {"answer": answer}

    def fallback(self, state: GraphState) -> GraphState:
        """Provides a fallback answer if documents are not relevant."""
        logger.info("Node: Providing fallback answer")
        answer = "I'm sorry, but the provided document does not contain information relevant to your question."
        return {"answer": answer}

    def create_rag_graph(self, vector_store):
        """Creates and compiles the RAG LangGraph."""
        try:
            if not self.llm: self.llm = self.load_llm()
            self.retriever = vector_store.as_retriever(search_kwargs={"k": 3})
            
            workflow = StateGraph(GraphState)
            workflow.add_node("retrieve", self.retrieve)
            workflow.add_node("generate", self.generate)
            workflow.add_node("fallback", self.fallback)

            workflow.set_entry_point("retrieve")
            workflow.add_conditional_edges(
                "retrieve",
                self.grade_documents,
                {"generate": "generate", "fallback": "fallback"}
            )
            workflow.add_edge("generate", END)
            workflow.add_edge("fallback", END)
            
            self.rag_graph = workflow.compile()
            logger.info("RAG graph created and compiled successfully")
        except Exception as e:
            logger.error(f"Error creating RAG graph: {e}")
            raise Exception(f"Failed to create RAG graph: {str(e)}")

    def answer_question(self, question: str) -> Dict[str, Any]:
        """Answer a question using the RAG graph."""
        try:
            if not self.rag_graph: raise Exception("RAG graph not initialized")
            result_state = self.rag_graph.invoke({"question": question})
            return {
                "answer": result_state["answer"],
                "source_documents": result_state.get("documents", [])
            }
        except Exception as e:
            logger.error(f"Error answering question with graph: {e}")
            raise Exception(f"Failed to answer question: {str(e)}")

# --- Helper Functions ---

def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[LangChainDocument]:
    """Split text into chunks and create LangChain documents."""
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap, length_function=len)
    chunks = text_splitter.split_text(text)
    documents = [LangChainDocument(page_content=chunk) for chunk in chunks]
    logger.info(f"Text split into {len(documents)} chunks")
    return documents

def get_file_hash(uploaded_file) -> str:
    """Generate hash for uploaded file to enable caching."""
    file_content = uploaded_file.read()
    uploaded_file.seek(0)
    return hashlib.md5(file_content).hexdigest()

def display_chat_message(role: str, content: str, context: Optional[List] = None):
    """Display a chat message with proper styling."""
    css_class = "user-message" if role == "user" else "bot-message"
    icon = "🧑" if role == "user" else "🤖"
    
    st.markdown(f'<div class="chat-message {css_class}"><strong>{icon} {role.title()}:</strong><br>{content}</div>', unsafe_allow_html=True)
    
    if context and st.session_state.get('show_context', False):
        with st.expander("📄 Source Context", expanded=False):
            for i, doc in enumerate(context[:2]):
                st.markdown(f'<div class="context-box"><strong>Source {i+1}:</strong><br>{doc.page_content[:500]}...</div>', unsafe_allow_html=True)

# --- Main Streamlit Application ---

def main():
    """Main Streamlit application."""
    if 'chat_history' not in st.session_state: st.session_state.chat_history = []
    if 'vector_store' not in st.session_state: st.session_state.vector_store = None
    if 'qa_system' not in st.session_state: st.session_state.qa_system = QASystem()
    if 'embedding_manager' not in st.session_state: st.session_state.embedding_manager = EmbeddingManager()
    if 'processed_file_hash' not in st.session_state: st.session_state.processed_file_hash = None
    if 'show_context' not in st.session_state: st.session_state.show_context = True
    
    st.title("📚 Document Q&A with LangGraph")
    st.markdown("Upload a document and ask questions. This system uses LangGraph to validate document relevance before answering.")
    
    with st.sidebar:
        st.header("📁 Document Upload")
        uploaded_file = st.file_uploader("Choose a file", type=['pdf', 'docx', 'txt'])
        st.header("⚙️ Settings")
        chunk_size = 1000
        chunk_overlap = 200
        st.session_state.show_context = st.checkbox("Show Source Context", value=True)
        if st.button("🗑️ Clear Chat History"):
            st.session_state.chat_history = []
            st.rerun()
        if st.button("🔄 Reset System"):
            st.session_state.clear()
            st.rerun()

    col1, col2 = st.columns([2, 1])

    with col1:
        if uploaded_file is not None:
            try:
                current_file_hash = get_file_hash(uploaded_file)
                if st.session_state.processed_file_hash != current_file_hash or st.session_state.vector_store is None:
                    with st.spinner("Processing document... This may take a moment."):
                        text = DocumentProcessor.process_uploaded_file(uploaded_file)
                        if not text.strip():
                            st.error("No text could be extracted from the file.")
                            return
                        documents = chunk_text(text, chunk_size, chunk_overlap)
                        st.session_state.vector_store = st.session_state.embedding_manager.create_vector_store(documents)
                        st.session_state.qa_system.create_rag_graph(st.session_state.vector_store)
                        st.session_state.processed_file_hash = current_file_hash
                    st.success(f"Document processed! Created {len(documents)} text chunks.")
            except Exception as e:
                st.markdown(f'<div class="error-message"><strong>❌ Error processing document:</strong><br>{e}</div>', unsafe_allow_html=True)
                return

        if st.session_state.vector_store is not None:
            st.subheader("💬 Ask a Question")
            with st.form(key="question_form", clear_on_submit=True):
                question = st.text_input("Enter your question:", placeholder="What is this document about?", key="question_input")
                ask_button = st.form_submit_button("🔍 Ask", type="primary")
            
            if ask_button and question.strip():
                try:
                    with st.spinner("Thinking..."):
                        result = st.session_state.qa_system.answer_question(question)
                        st.session_state.chat_history.append({"question": question, **result})
                        st.rerun()
                except Exception as e:
                    st.markdown(f'<div class="error-message"><strong>❌ Error answering question:</strong><br>{e}</div>', unsafe_allow_html=True)
        else:
            st.info("👆 Please upload a document to start the Q&A process.")
        
        if st.session_state.chat_history:
            st.subheader("💭 Conversation History")
            for i, chat in enumerate(reversed(st.session_state.chat_history)):
                display_chat_message("user", chat["question"])
                display_chat_message("assistant", chat["answer"], chat.get("source_documents"))
                if i < len(st.session_state.chat_history) - 1: st.divider()
    
    with col2:
        if uploaded_file is not None:
            st.subheader("📄 Document Info")
            st.write(f"**Filename:** {uploaded_file.name}")
            st.write(f"**File Size:** {uploaded_file.size/1024:.2f} KB")
            if st.session_state.vector_store:
                st.write(f"**Text Chunks:** {st.session_state.vector_store.index.ntotal}")
                st.write("**Status:** ✅ Ready for Q&A")
        
        st.subheader("🔧 System Info")
        st.write("**Graph:** LangGraph RAG")
        st.write("**LLM:** google/flan-t5-large")
        st.write("**Embeddings:** all-MiniLM-L6-v2")
        auth_status = "✅ Authenticated" if HF_TOKEN else "⚠️ No Token"
        st.write(f"**HF Auth:** {auth_status}")
        st.write(f"**Device:** {'CUDA' if torch.cuda.is_available() else 'CPU'}")

if __name__ == "__main__":
    main()
